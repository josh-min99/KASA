"""
한글(HWP)에서 바로 열리는 문서 파일을 만든다.

왜 .docx 인가
  HWP 5.0 은 OLE 복합문서 기반 폐쇄 포맷이라 외부에서 유효한 파일을 만들기 어렵다.
  한글 2010 이후는 .docx 를 그대로 열고, 열어서 '한글 문서(.hwp)로 저장' 하면 된다.
  표·그림·글꼴·줄간격이 모두 유지되므로 실무상 가장 확실한 경로다.

서식 (KASA 연구계획서 양식)
  □ 로 시작하는 줄 : HY헤드라인M 14pt
  ◦ / - / ※ 줄     : 휴먼명조 14pt, 줄간격 160%, 자간 0%
  그림 캡션         : 휴먼명조 12pt (본문보다 한 단계 작게)
  ※ 자간 0% 는 기본값이므로 별도 지정하지 않는다.

원본은 docs/한글양식_2_연구방법_3_예상결과.html 이다. 내용 수정은 HTML 에만 한다.

실행: python src/v3_make_hwp_docx.py
"""
import os
import re
import sys
from html.parser import HTMLParser

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(ROOT, "docs")
SRC = os.path.join(DOCS, "한글양식_2_연구방법_3_예상결과.html")
DST = os.path.join(DOCS, "KASA_연구계획서_2연구방법_3예상결과.docx")

F_HEAD = "HY헤드라인M"
F_BODY = "휴먼명조"
SZ_HEAD, SZ_BODY, SZ_CAP = 14, 14, 12
FIG_WIDTH_CM = 15.5          # A4 기본 여백에서 본문 폭에 맞춘 값

# 원고에는 팀 내부용 메모가 섞여 있다. 제출본에는 남으면 안 되므로 여기서 걸러낸다.
# (HTML 원본에는 그대로 두어야 편집할 때 판단 근거가 남는다.)
DROP_PARA = [
    "KASA 연구계획서 —",                      # 작업용 표제
    "드라이랩 담당분 원고",
    "아래는 '나. 드라이랩' 부분입니다",        # 붙이는 위치 안내
]
DROP_INLINE = [
    "[축약 가능]", "[여유 시 생략 가능 — 표 1 과 내용이 겹침]",
    "[여유 시 생략 가능]", "[필수 — 양식이 요구하는 시뮬레이션 자료]",
    "[팀 전체 절차에 아래 항목을 삽입]",
    "→ 예상 결과 (1), (2)", "→ 예상 결과 (2), (3)",
    "→ 활용방안 절로 이관 제안",
]


def strip_notes(text):
    for m in DROP_INLINE:
        text = text.replace(m, "")
    return re.sub(r"\s{2,}", " ", text).strip()


def set_font(run, name, size, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # 한글 글꼴은 eastAsia 속성까지 지정해야 적용된다
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)


def body_para(doc, text, indent_cm=0.0, size=SZ_BODY, font=F_BODY,
              space_after=2, bold_spans=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indent_cm)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.6                     # 줄간격 160%
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    for seg, is_bold in (bold_spans or [(text, False)]):
        set_font(p.add_run(seg), font, size, bold=is_bold)
    return p


class Doc(HTMLParser):
    """HTML 원고를 docx 로 옮긴다. 굵게 표시와 표·그림을 살린다."""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.buf = []            # [(텍스트, 굵기)]
        self.bold = 0
        self.skip = 0            # 안내 박스 등 계획서에 안 들어갈 블록
        self.in_table = False
        self.rows, self.row, self.cell = [], [], []
        self.in_cell = False
        self.in_pre = False
        self.pending_img = None

    # ---------------------------------------------------------------- 수집
    def handle_data(self, d):
        if self.skip:
            return
        if self.in_pre:
            self.emit_pre(d)
            return
        if self.in_cell:
            self.cell.append((d, self.bold > 0))
        else:
            self.buf.append((d, self.bold > 0))

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        if tag == "div" and "background:#fafafa" in a.get("style", ""):
            self.skip += 1                     # 작성 안내 박스는 제외
            return
        if self.skip:
            if tag == "div":
                self.skip += 1
            return
        if tag in ("b", "strong"):
            self.bold += 1
        elif tag == "br":
            (self.cell if self.in_cell else self.buf).append((" ", False))
        elif tag == "table":
            self.flush(); self.in_table, self.rows = True, []
        elif tag == "tr" and self.in_table:
            self.row = []
        elif tag in ("td", "th") and self.in_table:
            self.in_cell, self.cell = True, []
        elif tag == "pre":
            self.flush(); self.in_pre = True
        elif tag == "img":
            self.flush(); self.pending_img = a.get("src", "")
        elif tag == "hr":
            self.flush()

    def handle_endtag(self, tag):
        if self.skip:
            if tag == "div":
                self.skip -= 1
            return
        if tag in ("b", "strong"):
            self.bold = max(0, self.bold - 1)
        elif tag == "table":
            self.render_table(); self.in_table = False
        elif tag == "tr" and self.in_table:
            if self.row:
                self.rows.append(self.row)
        elif tag in ("td", "th") and self.in_table:
            txt = re.sub(r"\s+", " ", "".join(t for t, _ in self.cell)).strip()
            anyb = any(b for _, b in self.cell)
            self.row.append((txt, anyb))
            self.in_cell = False
        elif tag == "pre":
            self.in_pre = False
        elif tag == "p":
            self.flush()

    # ---------------------------------------------------------------- 출력
    def emit_pre(self, d):
        for line in d.strip("\n").splitlines():
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1.4)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.2
            pf.space_after = Pt(0)
            set_font(p.add_run(line), "함초롬돋움", 11)

    def flush(self):
        if self.pending_img:
            self.add_image(self.pending_img)
            self.pending_img = None
        if not self.buf:
            return
        spans = [(t, b) for t, b in self.buf if t.strip() or t == " "]
        self.buf = []
        text = "".join(t for t, _ in spans)
        norm = re.sub(r"\s+", " ", text).strip()
        if not norm:
            return
        if any(k in norm for k in DROP_PARA):
            return
        spans = [(strip_notes(t) if any(m in t for m in DROP_INLINE) else t, b)
                 for t, b in spans]
        spans = [(t, b) for t, b in spans if t]
        norm = strip_notes(norm)
        if not norm:
            return
        # 앞뒤 공백 정리하되 굵기 구간은 유지
        spans = [(re.sub(r"\s+", " ", t), b) for t, b in spans]
        if spans:
            spans[0] = (spans[0][0].lstrip(), spans[0][1])
            spans[-1] = (spans[-1][0].rstrip(), spans[-1][1])
        spans = [(t, b) for t, b in spans if t]

        if norm.startswith("2 ") or norm.startswith("3 "):     # 절 제목
            body_para(self.doc, "", size=SZ_HEAD)
            body_para(self.doc, norm, font=F_HEAD, size=SZ_HEAD,
                      bold_spans=[(norm, True)], space_after=6)
        elif norm.startswith("□"):
            body_para(self.doc, norm, font=F_HEAD, size=SZ_HEAD,
                      bold_spans=[(norm, True)], space_after=4)
        elif norm.startswith("◦"):
            body_para(self.doc, norm, indent_cm=0.0, bold_spans=spans, space_after=2)
        elif norm.startswith("-"):
            body_para(self.doc, norm, indent_cm=0.7, bold_spans=spans, space_after=2)
        elif norm.startswith("※"):
            body_para(self.doc, norm, indent_cm=0.7, size=SZ_CAP, bold_spans=spans)
        elif norm.startswith("[") or norm.startswith("그림"):
            body_para(self.doc, norm, indent_cm=0.7, size=SZ_CAP, bold_spans=spans,
                      space_after=8)
        else:
            body_para(self.doc, norm, indent_cm=0.4, bold_spans=spans)

    def add_image(self, rel):
        path = os.path.normpath(os.path.join(DOCS, rel))
        if not os.path.exists(path):
            print("  그림 없음:", rel)
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(path, width=Cm(FIG_WIDTH_CM))
        print(f"  그림 삽입: {os.path.basename(path)}")

    def render_table(self):
        rows = [r for r in self.rows if r]
        if not rows:
            return
        ncol = max(len(r) for r in rows)
        t = self.doc.add_table(rows=len(rows), cols=ncol)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, r in enumerate(rows):
            for j in range(ncol):
                txt, bold = r[j] if j < len(r) else ("", False)
                cell = t.cell(i, j)
                cell.text = ""
                para = cell.paragraphs[0]
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.15
                set_font(para.add_run(txt), F_BODY, 11, bold=(bold or i == 0))
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)


def main():
    html = open(SRC, encoding="utf-8").read()
    doc = Document()

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)      # A4
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.0)

    conv = Doc(doc)
    conv.feed(html)
    conv.flush()
    doc.save(DST)

    n_par = len(doc.paragraphs)
    print(f"\n-> {DST}")
    print(f"   문단 {n_par}개, 표 {len(doc.tables)}개")
    print("   한글에서 열어 '다른 이름으로 저장 > 한글 문서(*.hwp)' 하면 됩니다.")


if __name__ == "__main__":
    main()
