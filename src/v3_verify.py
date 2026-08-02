"""
R9: 문서에 쓴 수치를 산출물과 대조한다.

`docs/계획서_드라이랩_v3.html` 과 `docs/제안_타절_v3.md` 에 적은 숫자를 전건 검사한다.
문서를 고쳤는데 산출물을 다시 안 돌렸거나, 산출물이 바뀌었는데 문서를 안 고친 경우를 잡는다.

실행: python src/v3_verify.py     (전건 통과 시 exit 0)
"""
import os
import re
import sys

import numpy as np
import pandas as pd

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA = os.path.join(ROOT, "data", "rhythm")
RES = os.path.join(ROOT, "results", "v3")
DOCS = os.path.join(ROOT, "docs")

FAILS = []
OKS = []


def chk(name, got, want, tol=0.0):
    if isinstance(want, (int, float)) and isinstance(got, (int, float, np.floating, np.integer)):
        ok = abs(float(got) - float(want)) <= tol
    else:
        ok = (got == want)
    (OKS if ok else FAILS).append((name, got, want))
    print(f"  {'OK ' if ok else '!! '}{name:58s} 산출={got}  문서={want}")


def main():
    print("=" * 92)
    print("R9  문서 수치 대조")
    print("=" * 92)

    au = pd.read_csv(os.path.join(DATA, "audit.csv"))
    cand = pd.read_csv(os.path.join(DATA, "osdr_candidates.csv"))
    # 전 파일 색인은 20MB 라 git 에서 제외했다(src/v3_r1_osdr_sweep.py 로 재생성).
    # 없으면 감사 요약문에 기록된 수를 쓴다.
    idx_path = os.path.join(DATA, "osdr_files_index.csv")
    if os.path.exists(idx_path):
        n_files = int(len(pd.read_csv(idx_path)))
    else:
        s = open(os.path.join(RES, "r1_audit_summary.txt"), encoding="utf-8").read()
        m = re.search(r"파일\s+([\d,]+)개 전수", s)
        n_files = int(m.group(1).replace(",", "")) if m else -1
    long = pd.read_csv(os.path.join(DATA, "long.csv"))
    par = pd.read_csv(os.path.join(DATA, "params.csv"))
    co = pd.read_csv(os.path.join(DATA, "concordance.csv"))
    pr = pd.read_csv(os.path.join(DATA, "phase_precision.csv"))
    need = pd.read_csv(os.path.join(DATA, "power_required_n.csv"))
    inv = pd.read_csv(os.path.join(DATA, "inverse_dphi.csv"))

    print("\n[탐색 규모]")
    chk("OSDR 파일 색인 수", n_files, 157801)
    chk("OSDR 리듬 후보 파일 수", int(len(cand)), 52)
    chk("OSDR 리듬 후보 스터디 수", int(cand.OSD_ID.nunique()), 5)
    chk("공통 스키마 행 수", int(len(long)), 16561)

    print("\n[감사 판정 건수]")
    vc = au.verdict.value_counts()
    for v, want in [("PASS", 1), ("CONDITIONAL", 1), ("LEVEL_ONLY", 1),
                    ("DIRECTION_ONLY", 2), ("LITERATURE", 6), ("REJECT", 5)]:
        chk(f"판정 {v}", int(vc.get(v, 0)), want)
    chk("후보 총건수", int(len(au)), 16)

    print("\n[게이트 G2 — Helissen 재현]")
    a = par[(par.dataset == "helissen2020") & (par.variable == "activity")]
    chk("활동 MESOR baseline", round(float(a[a.condition == "baseline"].mesor.mean()), 3), 0.272, 0.001)
    chk("활동 MESOR HLU중", round(float(a[a.condition == "treatment"].mesor.mean()), 3), 0.046, 0.001)
    chk("활동 리듬검출 p baseline", round(float(a[a.condition == "baseline"].p_rhythm.mean()), 3), 0.088, 0.001)
    chk("활동 리듬검출 p HLU중", round(float(a[a.condition == "treatment"].p_rhythm.mean()), 3), 0.412, 0.001)

    print("\n[정합성 — 활동량 수준]")
    lv = co[(co.axis == "활동량 수준") & co.value.notna()]
    vals = {r.arm: r for _, r in lv.iterrows()}
    for arm, want in [("지상 언로딩 설치류 (피하 코호트 A)", -0.800),
                      ("지상 언로딩 설치류 (피하 코호트 B)", -0.793),
                      ("지상 언로딩 설치류 (심부 코호트)", -0.494)]:
        chk(f"활동량 수준 {arm[-8:]}", round(float(vals[arm].value), 3), want, 0.001)
    fly = vals["실제 비행 (중력만 분리, 초파리)"]
    chk("초파리 uG 대 1G 활동 수준", round(float(fly.value), 3), 0.027, 0.001)
    chk("초파리 CI 하한", round(float(fly.ci_lo), 3), -0.002, 0.002)
    chk("초파리 CI 상한", round(float(fly.ci_hi), 3), 0.060, 0.002)
    chk("지상 활동량 수준 3건 모두 0 배제",
        int(sum((r.ci_lo > 0) or (r.ci_hi < 0) for _, r in lv.iterrows()
                if "지상" in r.arm)), 3)

    print("\n[정합성 — 체온 acrophase]")
    tb = co[(co.axis == "체온 acrophase") & co.value.notna()]
    chk("체온 acrophase 비교 개수", int(len(tb)), 5)
    chk("점추정이 모두 지연(양수)", int((tb.value > 0).sum()), 5)
    chk("구간이 0을 배제하는 개수",
        int(sum((r.ci_lo > 0) or (r.ci_hi < 0) for _, r in tb.iterrows())), 3)

    print("\n[마스킹 — 위상 추정 정밀도]")
    s = pr.groupby(["variable", "dataset", "condition"]).phase_sd_h.median().reset_index()
    b = s[s.condition == "baseline"]
    t = s[s.condition == "treatment"]
    chk("baseline 위상 SD 최소(h)", round(float(b.phase_sd_h.min()), 2), 0.14, 0.005)
    chk("baseline 위상 SD 최대(h)", round(float(b.phase_sd_h.max()), 2), 1.11, 0.005)
    chk("HLU중 위상 SD 최소(h)", round(float(t.phase_sd_h.min()), 2), 2.77, 0.005)
    chk("HLU중 위상 SD 최대(h)", round(float(t.phase_sd_h.max()), 2), 5.22, 0.005)
    infl = []
    for _, r in b.iterrows():
        m = t[(t.variable == r.variable) & (t.dataset == r.dataset)]
        if len(m):
            infl.append(float(m.phase_sd_h.iloc[0]) / float(r.phase_sd_h))
    chk("정밀도 저하 최소 배수", round(min(infl), 1), 3.6, 0.06)
    chk("정밀도 저하 최대 배수", round(max(infl), 0), 38.0, 0.6)
    core = [float(t[(t.variable == 'tb_core')].phase_sd_h.iloc[0]) /
            float(b[(b.variable == 'tb_core')].phase_sd_h.iloc[0])]
    chk("심부체온 저하 배수", round(core[0], 1), 9.0, 0.06)

    print("\n[검정력 — 필요 마리수 (처치 14일)]")
    def nreq(var, dphi):
        v = need[(need.variable == var) & (need.dphi == dphi)].n_needed
        return sorted(set(int(x) for x in v.dropna())), int(v.isna().sum())
    for var, dphi, want_vals, want_na in [
            ("tb_core", 1.0, [8], 0), ("tb_core", 2.0, [4], 0), ("tb_core", 3.0, [4], 0),
            ("tb_sub", 1.0, [12, 16], 0), ("tb_sub", 2.0, [6], 0), ("tb_sub", 3.0, [4], 0),
            ("activity", 2.0, [6, 8], 0), ("activity", 3.0, [4, 6], 0)]:
        got, na = nreq(var, dphi)
        chk(f"{var} dphi={dphi} 필요 n", str(got), str(want_vals))
    got, na = nreq("activity", 1.0)
    chk("activity dphi=1 에 16마리로 부족한 코호트 존재", int(na >= 1), 1)

    print("\n[역문제]")
    fl = inv[inv.contrast == "우주비행 vs 지상대조"].iloc[0]
    chk("비행 대비 잡음 바닥", round(float(fl.noise_floor), 3), 0.088, 0.001)
    chk("비행 대비 실측", round(float(fl.observed), 3), 0.215, 0.001)
    chk("비행 대비 검출가능 Δφ 하한(h)", round(float(fl.dphi_min_detectable), 2), 0.55, 0.03)
    chk("실측과 정합적인 Δφ(h)", round(float(fl.dphi_matching_observed), 2), 1.30, 0.03)
    chk("전 대비 중 최소 Δφ 하한(h)", round(float(inv.dphi_min_detectable.min()), 2), 0.40, 0.03)

    print("\n[문서에 있어야 할 문자열]")
    html = open(os.path.join(DOCS, "계획서_드라이랩_v3.html"), encoding="utf-8").read()
    md = open(os.path.join(DOCS, "제안_타절_v3.md"), encoding="utf-8").read()
    for token in ["157,801", "No data submitted by PI", "AGBRESA", "0.055" if False else "0.55",
                  "16,561", "심부체온", "free-run", "잡음 대비"]:
        chk(f"계획서에 '{token}' 존재", token in html, True)
    for token in ["Neurolab STS-90", "OSD-595", "Hélissen", "Martin et al. 2020"]:
        chk(f"제안서에 '{token}' 존재", token in md, True)

    # KASA 한글 양식 원고 (구글 독스 붙여넣기용 HTML 과 txt 는 같은 내용이어야 한다)
    hwp_html = open(os.path.join(DOCS, "한글양식_2_연구방법_3_예상결과.html"),
                    encoding="utf-8").read()
    hwp_txt = open(os.path.join(DOCS, "한글양식_2_연구방법_3_예상결과.txt"),
                   encoding="utf-8").read()
    for token in ["157,801", "2,593", "633", "3.6배에서 38배", "0.55시간",
                  "No data submitted by PI", "자유진행(free-run)",
                  "표 1.", "표 2.", "그림 1.", "그림 2.",
                  "심부체온 5/6", "절편 검정", "기울기 검정",
                  "(T − 24)시간씩 이동"]:
        chk(f"한글양식 원고에 '{token}' 존재", token in hwp_html, True)
        chk(f"  (txt 판에도) '{token}'", token in hwp_txt, True)

    # 표 4 (군별 예상 판정 간격) 수치가 산출물과 맞는가
    print("\n[표 4 — 군별 예상 판정 간격]")
    pw = pd.read_csv(os.path.join(DATA, "predicted_wetlab.csv"))
    pw["ratio"] = pw.gap_h.abs() / pw.sem_free
    for T, gap, ratio in [(24, 4.2, 7.1), (20, 3.8, 6.4), (28, 11.8, 20.0)]:
        r = pw[pw.T_hours == T].iloc[0]
        chk(f"T={T}h 판정 간격(h)", round(abs(float(r.gap_h)), 1), gap, 0.05)
        chk(f"T={T}h 표준오차 대비 배수", round(float(r.ratio), 1), ratio, 0.05)
    chk("세 군 모두 검출 가능", int(pw.detectable.all()), 1)
    # T=28 이 T=20 보다 몇 배 강한 자극을 요구하는가 (문서 기재값 2.6배)

    print("\n[그림 첨부 상태]")
    figs = ["그림2_마스킹과검정력.png", "그림3_웻랩예상결과.png"]
    figdir = os.path.join(RES, "figures")
    for f in figs:
        chk(f"{f} 존재", os.path.exists(os.path.join(figdir, f)), True)
        chk(f"  본문에 삽입됨", f in hwp_html, True)
    built = os.path.join(DOCS, "한글양식_최종_이미지포함.html")
    chk("이미지 포함 배포본 존재", os.path.exists(built), True)
    if os.path.exists(built):
        b = open(built, encoding="utf-8").read()
        chk("  배포본에 그림이 base64 로 박혀 있음",
            b.count("data:image/png;base64,"), len(figs))
        chk("  배포본에 남은 상대경로 없음", ".png\"" in b, False)

    print("\n[리듬 유지 개체 수 — 심부체온 대 활동량]")
    for ds, var, want_keep, want_tot in [("helissen2022", "tb_core", 5, 6),
                                         ("helissen2022", "activity", 1, 6),
                                         ("helissen2020", "tb_sub", 1, 5)]:
        t = par[(par.dataset == ds) & (par.variable == var) &
                (par.condition == "treatment")]
        chk(f"{var} HLU중 리듬 유지 개체", f"{int((t.p_rhythm < 0.05).sum())}/{len(t)}",
            f"{want_keep}/{want_tot}")

    check_docx()

    print("\n" + "=" * 92)
    print(f"통과 {len(OKS)} / 실패 {len(FAILS)}")
    if FAILS:
        print("\n불일치:")
        for n, g, w in FAILS:
            print(f"  {n}: 산출={g} 문서={w}")
    return 1 if FAILS else 0


def check_docx():
    """한글 제출본(.docx) 점검 — 표·그림 개수와 작업 메모 잔류 여부."""
    print("\n[한글 제출본]")
    dx = os.path.join(DOCS, "KASA_연구계획서_2연구방법_3예상결과.docx")
    chk("docx 생성됨", os.path.exists(dx), True)
    if not os.path.exists(dx):
        return
    from docx import Document
    dd = Document(dx)
    body = "\n".join(p.text for p in dd.paragraphs)
    chk("  표 2개", len(dd.tables), 2)
    n_img = sum(1 for r in dd.part.rels.values() if "image" in r.reltype)
    chk("  그림 2개", n_img, 2)
    for note in ["축약 가능", "여유 시 생략", "팀 전체 절차", "이관 제안",
                 "드라이랩 담당분"]:
        chk(f"  작업 메모 '{note}' 없음", note in body, False)


if __name__ == "__main__":
    sys.exit(main())
